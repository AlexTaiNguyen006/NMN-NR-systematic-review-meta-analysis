#!/usr/bin/env python3
"""
Generate manuscript-ready supplementary .docx files:
  1. Supplementary_Tables.docx  - all supplementary tables S1-S10
  2. Supplementary_Figures.docx - all supplementary figures S1-S7
"""

import csv
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# Note: Since the script was moved to Revisions/, we step back one folder to match original structure
ROOT_DIR = os.path.dirname(BASE_DIR)
TABLES_DIR = os.path.join(ROOT_DIR, "results", "tables")
FIGURES_DIR = os.path.join(ROOT_DIR, "results", "figures")
OUTPUT_DIR = os.path.join(BASE_DIR, "submission_revised", "Supplementary_Material")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def set_cell_font(cell, size=8, bold=False):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.name = "Times New Roman"
            run.bold = bold
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shading_elem)


def add_table_from_csv(doc, csv_path, font_size=7.5, landscape=False):
    """Read a CSV and insert it as a formatted Word table."""
    with open(csv_path, "r", encoding="utf-8") as f:
        reader = csv.reader(f)
        rows = list(reader)

    if not rows:
        return

    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for j, val in enumerate(rows[0]):
        cell = table.cell(0, j)
        cell.text = val.strip()
        set_cell_font(cell, size=font_size, bold=True)
        set_cell_shading(cell, "D9E2F3")  # light blue header

    # Data rows
    for i, row in enumerate(rows[1:], start=1):
        for j, val in enumerate(row):
            if j < n_cols:
                cell = table.cell(i, j)
                cell.text = val.strip()
                set_cell_font(cell, size=font_size)
                # alternating row shading
                if i % 2 == 0:
                    set_cell_shading(cell, "F2F2F2")


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.name = "Times New Roman"
    run.italic = True
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)


def add_section_landscape(doc):
    """Add a new landscape section."""
    new_section = doc.add_section()
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Cm(29.7)
    new_section.page_height = Cm(21.0)
    new_section.left_margin = Cm(1.5)
    new_section.right_margin = Cm(1.5)
    new_section.top_margin = Cm(1.5)
    new_section.bottom_margin = Cm(1.5)
    return new_section


def add_section_portrait(doc):
    """Add a new portrait section."""
    new_section = doc.add_section()
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width = Cm(21.0)
    new_section.page_height = Cm(29.7)
    new_section.left_margin = Cm(2.54)
    new_section.right_margin = Cm(2.54)
    new_section.top_margin = Cm(2.54)
    new_section.bottom_margin = Cm(2.54)
    return new_section


# ---------------------------------------------------------------------------
# Table definitions: (file, title, note, landscape?)
# ---------------------------------------------------------------------------

SUPP_TABLES = [
    (
        "supp_table_S1_search_strategy.csv",
        "Table S1. Search Strategy by Database",
        "Searches conducted on 6 March 2026 across five databases. Duplicates removed via DOI, PMID, and fuzzy title matching.",
        True,
    ),
    (
        "supp_table_S2_study_characteristics.csv",
        "Table S2. Characteristics of Included Studies",
        "NMN, nicotinamide mononucleotide; NR, nicotinamide riboside; RCT, randomized controlled trial; "
        "BMI, body mass index; PD, Parkinson's disease; CKD, chronic kidney disease; COPD, chronic obstructive pulmonary disease; "
        "MCI, mild cognitive impairment; RoB, risk of bias.",
        True,
    ),
    (
        "supp_table_S3_full_extraction.csv",
        "Table S3. Full Data Extraction for All Included Studies and Outcomes",
        "MD, mean difference; SE, standard error; SD, standard deviation; SEM, standard error of the mean. "
        "Values represent post-intervention group means. Crossover SDs estimated from reported statistics where necessary.",
        True,
    ),
    (
        "supp_table_S4_rob2_detailed.csv",
        "Table S4. Risk of Bias 2 (RoB 2) Domain-Level Assessments",
        "D1, bias from the randomization process; D2, bias due to deviations from intended interventions; "
        "D3, bias due to missing outcome data; D4, bias in measurement of the outcome; D5, bias in selection of the reported result. "
        "Assessments performed independently by two reviewers with disagreements resolved by discussion.",
        True,
    ),
    (
        "supp_table_S5_pairwise_all.csv",
        "Table S5. Pairwise Meta-Analysis Results for All Outcomes",
        "MD, mean difference; CI, confidence interval; k, number of studies. Random-effects model (DerSimonian-Laird); "
        "fixed-effect model (inverse-variance). Heterogeneity quantified by I-squared and tau-squared.",
        True,
    ),
    (
        "supp_table_S6_nma_all.csv",
        "Table S6. Network Meta-Analysis Results (Direct, Indirect, and Network Estimates)",
        "MD, mean difference; CI, confidence interval; k, number of contributing studies. "
        "Indirect comparisons estimated via the Bucher method. "
        "Network estimates combine direct and indirect evidence where applicable.",
        True,
    ),
    (
        "supp_table_S7_loo_summary.csv",
        "Table S7. Leave-One-Out Sensitivity Analysis Summary",
        "LOO, leave-one-out; MD, mean difference. Direction change indicates the pooled MD changed sign upon study omission. "
        "Significance change indicates the pooled p-value crossed 0.05. Analyses with k < 3 not performed.",
        True,
    ),
    (
        "supp_table_S8_grade_summary.csv",
        "Table S8. GRADE/CINeMA Certainty of Evidence Assessment for Indirect Comparisons (NMN vs NR)",
        "GRADE, Grading of Recommendations Assessment, Development and Evaluation; CINeMA, Confidence in Network Meta-Analysis. "
        "All indirect comparisons started at low certainty (observational-equivalent for indirect evidence) and were further downgraded "
        "based on within-study bias, reporting bias, indirectness, imprecision, heterogeneity, and incoherence.",
        True,
    ),
    (
        "supp_table_S9_inter_rater_agreement.csv",
        "Table S9. Inter-Rater Agreement Across Review Phases",
        "Cohen's kappa interpreted as: >0.80 almost perfect, 0.61-0.80 substantial, 0.41-0.60 moderate. "
        "FAD, final arbiter decision. All disagreements resolved by discussion between reviewers or third-party adjudication.",
        False,
    ),
    (
        "supp_table_S10_publication_bias.csv",
        "Table S10. Publication Bias Assessment (Funnel Plot Asymmetry and Egger's Test)",
        "Egger's regression test for funnel plot asymmetry. Tests with k < 3 have insufficient power and are not interpretable. "
        "Funnel plots generated for all comparisons with k >= 3.",
        True,
    ),
    (
        "supp_table_S11_harmonization.csv",
        "Table S11. Outcome Harmonization Matrix",
        "Summary of data extraction methodology, standard deviation (SD)/standard error (SE) handling, "
        "units, and crossover wash-out assumptions to harmonize outcomes for indirect comparison.",
        True,
    ),
    (
        "supp_table_S12_protocol_deviations.csv",
        "Table S12. Protocol Deviations",
        "Details of explicitly documented deviations between the initial PROSPERO registration and "
        "the final executed analysis.",
        True,
    ),
]


# ---------------------------------------------------------------------------
# Figure definitions: (png_name, label, caption)
# ---------------------------------------------------------------------------

SUPP_FIGURES = [
    # S1: Glycaemic outcomes
    (
        "forest_FBG_NMN_vs_PBO.png",
        "Figure S1a",
        "Forest plot of fasting blood glucose (FBG): NMN vs placebo. Random-effects model (DerSimonian-Laird). "
        "Squares represent individual study mean differences (MD) with 95% confidence intervals (CI); diamond represents the pooled estimate.",
    ),
    (
        "forest_FBG_NR_vs_PBO.png",
        "Figure S1b",
        "Forest plot of fasting blood glucose (FBG): NR vs placebo.",
    ),
    (
        "forest_HbA1c_NMN_vs_PBO.png",
        "Figure S1c",
        "Forest plot of glycated haemoglobin (HbA1c): NMN vs placebo.",
    ),
    (
        "forest_fasting_insulin_NMN_vs_PBO.png",
        "Figure S1d",
        "Forest plot of fasting insulin: NMN vs placebo.",
    ),
    # S2: Lipid outcomes
    (
        "forest_TC_NMN_vs_PBO.png",
        "Figure S2a",
        "Forest plot of total cholesterol (TC): NMN vs placebo.",
    ),
    (
        "forest_TC_NR_vs_PBO.png",
        "Figure S2b",
        "Forest plot of total cholesterol (TC): NR vs placebo.",
    ),
    (
        "forest_LDL_NMN_vs_PBO.png",
        "Figure S2c",
        "Forest plot of low-density lipoprotein cholesterol (LDL-C): NMN vs placebo.",
    ),
    (
        "forest_LDL_NR_vs_PBO.png",
        "Figure S2d",
        "Forest plot of low-density lipoprotein cholesterol (LDL-C): NR vs placebo.",
    ),
    (
        "forest_HDL_NMN_vs_PBO.png",
        "Figure S2e",
        "Forest plot of high-density lipoprotein cholesterol (HDL-C): NMN vs placebo.",
    ),
    (
        "forest_HDL_NR_vs_PBO.png",
        "Figure S2f",
        "Forest plot of high-density lipoprotein cholesterol (HDL-C): NR vs placebo.",
    ),
    (
        "forest_TG_NMN_vs_PBO.png",
        "Figure S2g",
        "Forest plot of triglycerides (TG): NMN vs placebo.",
    ),
    (
        "forest_TG_NR_vs_PBO.png",
        "Figure S2h",
        "Forest plot of triglycerides (TG): NR vs placebo.",
    ),
    # S3: Hepatic outcomes
    (
        "forest_ALT_NMN_vs_PBO.png",
        "Figure S3a",
        "Forest plot of alanine aminotransferase (ALT): NMN vs placebo.",
    ),
    (
        "forest_ALT_NR_vs_PBO.png",
        "Figure S3b",
        "Forest plot of alanine aminotransferase (ALT): NR vs placebo.",
    ),
    (
        "forest_AST_NMN_vs_PBO.png",
        "Figure S3c",
        "Forest plot of aspartate aminotransferase (AST): NMN vs placebo.",
    ),
    # S4: Blood pressure outcomes
    (
        "forest_SBP_NMN_vs_PBO.png",
        "Figure S4a",
        "Forest plot of systolic blood pressure (SBP): NMN vs placebo.",
    ),
    (
        "forest_DBP_NMN_vs_PBO.png",
        "Figure S4b",
        "Forest plot of diastolic blood pressure (DBP): NMN vs placebo.",
    ),
    # S5: Anthropometric outcomes
    (
        "forest_BMI_NMN_vs_PBO.png",
        "Figure S5a",
        "Forest plot of body mass index (BMI): NMN vs placebo.",
    ),
    (
        "forest_body_weight_NMN_vs_PBO.png",
        "Figure S5b",
        "Forest plot of body weight: NMN vs placebo.",
    ),
    # S6: Outcome reporting matrix
    (
        "outcome_reporting_matrix.png",
        "Figure S6",
        "Outcome reporting matrix showing which outcomes were reported (filled) or not reported (empty) across included studies. "
        "Rows represent metabolic outcomes; columns represent individual studies grouped by precursor (NMN or NR).",
    ),
    # S7: Leave-one-out sensitivity
    (
        "sensitivity_loo_forest_ALT_NMN_pairwise.png",
        "Figure S7a",
        "Leave-one-out sensitivity analysis forest plot for ALT (NMN vs placebo). Each row shows the pooled estimate after excluding the named study.",
    ),
    (
        "sensitivity_loo_forest_SBP_NMN_pairwise.png",
        "Figure S7b",
        "Leave-one-out sensitivity analysis forest plot for SBP (NMN vs placebo). Each row shows the pooled estimate after excluding the named study.",
    ),
    (
        "sensitivity_loo_forest_fasting_insulin_NMN_pairwise.png",
        "Figure S7c",
        "Leave-one-out sensitivity analysis forest plot for fasting insulin (NMN vs placebo). Each row shows the pooled estimate after excluding the named study.",
    ),
]


# ===========================================================================
# 1. Create Supplementary Tables document
# ===========================================================================

def create_tables_docx():
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)

    # --- Title page ---
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Supplementary Tables")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Nicotinamide Mononucleotide vs Nicotinamide Riboside Supplementation on "
        "Cardiometabolic Outcomes in Adults: A Systematic Review and Network Meta-Analysis"
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()
    toc_note = doc.add_paragraph()
    toc_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = toc_note.add_run("Contents:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    for fname, table_title, note, landscape in SUPP_TABLES:
        p = doc.add_paragraph()
        run = p.add_run(table_title)
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

    doc.add_page_break()

    # --- Individual tables ---
    for idx, (fname, table_title, note, landscape) in enumerate(SUPP_TABLES):
        csv_path = os.path.join(TABLES_DIR, fname)
        if not os.path.exists(csv_path):
            print(f"WARNING: {csv_path} not found, skipping.")
            continue

        # Switch to landscape for wide tables
        if landscape:
            add_section_landscape(doc)
        else:
            if idx > 0:
                add_section_portrait(doc)

        add_heading_styled(doc, table_title, level=2)
        add_table_from_csv(doc, csv_path, font_size=7)
        add_note(doc, f"Note: {note}")

        # page break after each table except the last
        if idx < len(SUPP_TABLES) - 1:
            pass  # section breaks handle pagination

    out_path = os.path.join(OUTPUT_DIR, "Supplementary_Tables.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


# ===========================================================================
# 2. Create Supplementary Figures document
# ===========================================================================

def create_figures_docx():
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)

    # --- Title page ---
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Supplementary Figures")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Nicotinamide Mononucleotide vs Nicotinamide Riboside Supplementation on "
        "Cardiometabolic Outcomes in Adults: A Systematic Review and Network Meta-Analysis"
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()
    toc_note = doc.add_paragraph()
    toc_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = toc_note.add_run("Contents:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    # Group figures by S-number for TOC
    current_group = None
    for png_name, label, caption in SUPP_FIGURES:
        # label is e.g. "Figure S1a" or "Figure S6"
        s_num = label.split()[-1]  # S1a, S2a, S6, etc.
        s_base = "".join(c for c in s_num if not c.islower())  # S1, S2, S6
        if s_base != current_group:
            current_group = s_base
            p = doc.add_paragraph()
            run = p.add_run(f"Figure {s_base}: See panels below")
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"

    doc.add_page_break()

    # Common note for forest plots
    forest_note = (
        "Random-effects model (DerSimonian-Laird). Squares represent individual study "
        "mean differences (MD) with 95% confidence intervals (CI); the diamond represents "
        "the pooled estimate. Heterogeneity assessed via I-squared and Cochran's Q."
    )

    # --- Individual figures ---
    for idx, (png_name, label, caption) in enumerate(SUPP_FIGURES):
        png_path = os.path.join(FIGURES_DIR, png_name)
        if not os.path.exists(png_path):
            print(f"WARNING: {png_path} not found, skipping.")
            continue

        # Figure label
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = heading.add_run(f"{label}.")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        run2 = heading.add_run(f" {caption}")
        run2.font.size = Pt(10)
        run2.font.name = "Times New Roman"

        # Insert image - full width for landscape-style figures
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(png_path, width=Inches(6.0))

        # Page break after each figure
        if idx < len(SUPP_FIGURES) - 1:
            doc.add_page_break()

    out_path = os.path.join(OUTPUT_DIR, "Supplementary_Figures.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


# ===========================================================================
# Main
# ===========================================================================

if __name__ == "__main__":
    print("Generating Supplementary Tables .docx ...")
    create_tables_docx()
    print()
    print("Generating Supplementary Figures .docx ...")
    create_figures_docx()
    print()
    print("Done.")
