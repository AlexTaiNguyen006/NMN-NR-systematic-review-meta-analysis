import docx

def dump_docx(file_path):
    doc = docx.Document(file_path)
    print("=== PARAGRAPHS ===")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"[{i}] {p.text}")
    print("\n=== TABLES ===")
    for i, t in enumerate(doc.tables):
        print(f"Table {i+1}:")
        for row in t.rows:
            print(" | ".join([c.text.replace('\n', ' ').strip() for c in row.cells]))
        print("-" * 40)

dump_docx('Revisions/submission_revised/NMN_NR_Systematic_Review_JNS.docx')
