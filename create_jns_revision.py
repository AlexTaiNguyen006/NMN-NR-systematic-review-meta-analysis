#!/usr/bin/env python3
"""
Create Public Health Nutrition submission files:
  1. Cover letter (txt) - for ScholarOne text box
  2. Title page (docx) - separate for double-blind review
  3. Trimmed manuscript (docx) - under 5000 words body text, no author info
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# ============================================================
# Directory setup
# ============================================================
BASE = 'submission_PHN'
os.makedirs(BASE, exist_ok=True)

# ============================================================
# 1. COVER LETTER
# ============================================================
cover_letter = """Cover Letter

Dear Editor-in-Chief,

We are pleased to submit our manuscript entitled "Mapping the Evidence Gap Between NMN and NR for Metabolic Outcomes: A Systematic Review, Transitivity Assessment, and Indirect Comparison Meta-Analysis" for consideration as a Systematic Review Article in Public Health Nutrition.

This systematic review addresses a critical gap in the NAD+ precursor supplementation literature. Nicotinamide mononucleotide (NMN) and nicotinamide riboside (NR) are widely marketed dietary supplements promoted for cardiometabolic health, yet no study has directly compared them in humans. Using the Bucher indirect comparison method across 8 RCTs (4 NMN, 4 NR; 73 effect sizes), we provide the first quantitative comparison for 14 cardiometabolic outcomes.

We found no statistically significant difference between NMN and NR for any outcome, with all evidence rated Very Low certainty (GRADE/CINeMA) due to imprecision and within-study bias. These findings are directly relevant to public health nutrition: consumer demand for NAD+ precursors is growing rapidly worldwide, and policymakers, clinicians, and consumers need rigorous evidence to guide supplementation decisions. Our results highlight that claims of superiority for either supplement currently lack evidentiary support and that adequately powered head-to-head trials are urgently needed.

The study follows PRISMA 2020 guidelines and was prospectively registered on PROSPERO (CRD420261330487). Dual-reviewer processes were applied throughout (screening, extraction, RoB 2 assessment). All data and analysis code are publicly available on GitHub and archived on Zenodo (DOI: 10.5281/zenodo.19403850). A preprint has been posted on bioRxiv (BIORXIV/2026/716917).

We believe this work aligns with PHN's scope of evaluating nutritional interventions and informing evidence-based nutrition practice, with relevance to an international readership given the global market for these supplements.

Sincerely,

Alexander Tai Nguyen (corresponding author)
Johns Hopkins University, Baltimore, MD, United States
tnguy277@jh.edu

Bryan Nguyen
University of Richmond, Richmond, VA, United States
"""

with open(os.path.join(BASE, 'cover_letter_PHN.txt'), 'w') as f:
    f.write(cover_letter.strip())
print("Created: cover_letter_PHN.txt")


# ============================================================
# 2. TITLE PAGE (separate docx for double-blind)
# ============================================================
def create_title_page():
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Mapping the Evidence Gap Between NMN and NR for Metabolic Outcomes:\n'
                     'A Systematic Review, Transitivity Assessment, and Indirect Comparison Meta-Analysis')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Alexander Tai Nguyen')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    sup = p.add_run('1')
    sup.font.superscript = True
    sup.font.name = 'Times New Roman'
    p.add_run('*, Bryan Nguyen').font.name = 'Times New Roman'
    sup2 = p.add_run('2')
    sup2.font.superscript = True
    sup2.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Affiliations
    doc.add_paragraph(
        '1 Johns Hopkins University, Baltimore, MD, United States\n'
        '2 University of Richmond, Richmond, VA, United States'
    )

    doc.add_paragraph()

    # Corresponding author
    p = doc.add_paragraph()
    run = p.add_run('*Corresponding author: ')
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run('Alexander Tai Nguyen, Johns Hopkins University, Baltimore, MD, United States. '
              'Email: tnguy277@jh.edu').font.name = 'Times New Roman'

    doc.add_paragraph()

    # Running title
    p = doc.add_paragraph()
    run = p.add_run('Running title: ')
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run('NMN vs NR Indirect Comparison MA').font.name = 'Times New Roman'
    # 33 characters - well under 45

    doc.add_paragraph()

    # Financial support
    p = doc.add_paragraph()
    run = p.add_run('Financial support: ')
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run('This research received no specific grant from any funding agency, '
              'commercial or not-for-profit sectors.').font.name = 'Times New Roman'

    doc.add_paragraph()

    # Declaration of interest
    p = doc.add_paragraph()
    run = p.add_run('Declaration of interest: ')
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run('None.').font.name = 'Times New Roman'

    doc.add_paragraph()

    # Authorship (CRediT)
    p = doc.add_paragraph()
    run = p.add_run('Authorship: ')
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run(
        'A.T.N.: Conceptualization, Data curation, Formal analysis, Investigation, '
        'Methodology, Project administration, Software, Validation, Visualization, '
        'Writing - original draft, Writing - review & editing. '
        'B.N.: Data curation, Investigation, Validation, Writing - review & editing.'
    ).font.name = 'Times New Roman'

    doc.add_paragraph()

    # Ethical standards disclosure
    p = doc.add_paragraph()
    run = p.add_run('Ethical standards disclosure: ')
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run('Not applicable. This is a systematic review of published studies; '
              'no primary research involving human participants was conducted.').font.name = 'Times New Roman'

    doc.add_paragraph()

    # Data availability
    p = doc.add_paragraph()
    run = p.add_run('Data availability: ')
    run.bold = True
    run.font.name = 'Times New Roman'
    p.add_run(
        'The complete dataset, analytical code, and all figures are publicly available at '
        'https://github.com/AlexTaiNguyen006/NMN-NR-systematic-review-meta-analysis '
        'and archived on Zenodo (DOI: 10.5281/zenodo.19403850).'
    ).font.name = 'Times New Roman'

    path = os.path.join(BASE, 'title_page_PHN.docx')
    doc.save(path)
    print("Created: title_page_PHN.docx")


# ============================================================
# 3. MANUSCRIPT (trimmed, double-blind, structured abstract)
# ============================================================
def create_manuscript():
    doc = Document()

    # Set margins and font
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # Configure heading styles
    for level in ['Heading 1', 'Heading 2', 'Heading 3']:
        if level in doc.styles:
            h = doc.styles[level]
            h.font.name = 'Times New Roman'
            h.font.color.rgb = RGBColor(0, 0, 0)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
        return h

    def add_para(text):
        p = doc.add_paragraph(text)
        return p

    def add_bold_inline(para, bold_text, normal_text):
        run = para.add_run(bold_text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run2 = para.add_run(normal_text)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)

    # ---- TITLE ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Mapping the Evidence Gap Between NMN and NR for Metabolic Outcomes:\n'
        'A Systematic Review, Transitivity Assessment, and Indirect Comparison Meta-Analysis'
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # blank line

    # ---- ABSTRACT (structured per PHN: Objective, Design, Setting, Participants, Results, Conclusions) ----
    add_heading('Abstract', level=1)

    p = doc.add_paragraph()
    add_bold_inline(p, 'Objective: ',
        'To characterize structural heterogeneity between nicotinamide mononucleotide (NMN) and '
        'nicotinamide riboside (NR) trial evidence bases, assess transitivity, and estimate indirect '
        'NMN versus NR effects for metabolic outcomes.')

    p = doc.add_paragraph()
    add_bold_inline(p, 'Design: ',
        'Systematic review with indirect comparison meta-analysis (Bucher method) following PRISMA 2020 '
        'guidelines. Risk of bias assessed using RoB 2; certainty of evidence evaluated using GRADE/CINeMA. '
        'Registered prospectively on PROSPERO (CRD420261330487).')

    p = doc.add_paragraph()
    add_bold_inline(p, 'Setting: ',
        'Five electronic databases searched from January 2018 to May 2025.')

    p = doc.add_paragraph()
    add_bold_inline(p, 'Participants: ',
        'Adults from fifteen RCTs (5 NMN, 10 NR; 740 participants) of oral NMN or NR versus placebo '
        'reporting metabolic outcomes.')

    p = doc.add_paragraph()
    add_bold_inline(p, 'Results: ',
        'Evidence bases were systematically asymmetric: NR was dosed 1.9-9.2 times higher than NMN on a '
        'molar basis; NMN trials enrolled predominantly East Asian populations while NR trials enrolled '
        'predominantly Western populations; and NAD+ assays were incompatible across precursors. Eight '
        'studies (4 NMN, 4 NR; 73 effect sizes) contributed to quantitative synthesis across 14 indirectly '
        'comparable outcomes. No pairwise meta-analysis reached significance for either precursor versus '
        'placebo for most outcomes. No indirect comparison reached statistical significance, and all were '
        'rated Very Low certainty. Leave-one-out sensitivity analyses showed no pairwise significance '
        'changes; one indirect change did not survive correction for multiple testing.')

    p = doc.add_paragraph()
    add_bold_inline(p, 'Conclusions: ',
        'Current evidence is structurally insufficient for reliable indirect comparison of NMN and NR. '
        'Barriers are quantifiable: future head-to-head trials should use equimolar dosing, harmonized '
        'NAD+ assays, minimum 24-week duration, and metabolically at-risk populations.')

    # Keywords
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_bold_inline(p, 'Keywords: ',
        'nicotinamide mononucleotide; nicotinamide riboside; NAD+; systematic review; '
        'indirect comparison; metabolic outcomes')

    p = doc.add_paragraph()
    add_bold_inline(p, 'Registration: ',
        'PROSPERO 2026 CRD420261330487; registered prior to data screening.')

    # ================================================================
    # 1. INTRODUCTION (~800 words)
    # ================================================================
    add_heading('1. Introduction', level=1)

    add_para(
        'Nicotinamide adenine dinucleotide (NAD+) is an essential coenzyme present in every '
        'living cell, participating in more than 500 enzymatic reactions spanning glycolysis, '
        'the tricarboxylic acid cycle, fatty acid oxidation, and mitochondrial electron '
        'transport(31). Beyond its role as a redox shuttle, NAD+ serves as a consumed substrate '
        'for three major families of signaling enzymes: the sirtuins (SIRT1-SIRT7), which '
        'regulate gene expression, mitochondrial biogenesis, and circadian rhythms(33); '
        'poly(ADP-ribose) polymerases (PARPs), which coordinate DNA damage repair; and the '
        'cyclic ADP-ribose hydrolase CD38, a major NAD+ consumer whose expression increases '
        'with age and chronic inflammation(32). Intracellular NAD+ levels decline progressively '
        'with age across multiple tissues, including skeletal muscle, liver, and adipose '
        'tissue(31,32). In murine models, this decline has been causally linked to insulin '
        'resistance, hepatic steatosis, endothelial dysfunction, and reduced exercise '
        'capacity(28,29). The concept of restoring NAD+ to youthful levels as a therapeutic '
        'strategy for metabolic and age-related diseases has therefore gained considerable '
        'scientific and commercial attention.'
    )

    add_para(
        'Two oral NAD+ precursor supplements have emerged as leading candidates for NAD+ '
        'augmentation in humans: nicotinamide mononucleotide (NMN) and nicotinamide riboside '
        '(NR)(1). NMN is a nucleotide composed of nicotinamide, ribose, and a phosphate group, '
        'and is a direct biosynthetic intermediate in the NAD+ salvage pathway. Its conversion '
        'to NAD+ requires a single adenylylation step catalyzed by nicotinamide mononucleotide '
        'adenylyltransferases (NMNAT1-3)(1). NR is a nucleoside (nicotinamide plus ribose '
        'without the phosphate group) that enters the salvage pathway one step earlier: it must '
        'first be phosphorylated by nicotinamide riboside kinases (NRK1/NRK2) to form NMN '
        'before undergoing the same NMNAT-catalyzed conversion(1). This biochemical difference '
        'has led to speculation that NMN may be more efficient at raising NAD+ levels because '
        'it bypasses the rate-limiting NRK phosphorylation step.'
    )

    add_para(
        'The pharmacokinetics of oral delivery add further complexity. NR is readily absorbed '
        'in the small intestine, but a substantial fraction is cleaved to nicotinamide by '
        'purine nucleoside phosphorylase before reaching the systemic circulation(30). NMN '
        'absorption was initially thought to require extracellular dephosphorylation to NR '
        'followed by re-phosphorylation inside the cell, but the identification of a putative '
        'NMN-specific transporter, Slc12a8, in murine gut epithelium suggests direct cellular '
        'uptake may occur(27). Whether Slc12a8-mediated transport operates efficiently in '
        'humans remains debated. These pharmacokinetic differences, combined with variations '
        'in supplement formulation and first-pass hepatic metabolism, may influence '
        'bioavailability independently of the intrinsic biochemistry.'
    )

    add_para(
        'As of May 2025, at least five NMN(2-6) and ten NR(7-16) RCTs have reported metabolic '
        'outcomes including glycemic markers, lipid profiles, hepatic enzymes, blood pressure, '
        'anthropometric measures, and circulating NAD+ concentrations. However, most trials are '
        'small (median n approximately 40), of short duration (8-12 weeks), and compare a single '
        'precursor against placebo. Critically, the two evidence bases diverge systematically: NMN '
        'trials have predominantly enrolled Asian populations (Japan, India) at doses of 250-300 '
        'mg/day, while NR trials have typically studied Western populations (Europe, North America) '
        'at 500-2,000 mg/day(2-16). No head-to-head RCT has directly compared NMN with NR on any '
        'clinical endpoint. These supplements are marketed globally with billions in annual sales, '
        'yet consumers and practitioners lack comparative evidence to guide choice between them - '
        'a question of direct relevance to public health nutrition practice.'
    )

    add_para(
        'Several recent systematic reviews have synthesized pairwise evidence for NMN or NR '
        'individually. Zheng et al. (2024) conducted the most comprehensive NMN-focused '
        'meta-analysis, pooling nine RCTs and reporting non-significant effects on most '
        'cardiometabolic outcomes(24). Alegre and Pastore (2024) reviewed both precursors '
        'narratively and concluded that clinical translation remained limited(25). Nascimento '
        'and Nogueira-de-Almeida (2024) focused on NR in cardiometabolic disease and found '
        'insufficient evidence to recommend clinical use(26). Importantly, none attempted to '
        'compare NMN against NR, because no closed network of direct evidence exists. However, '
        'network meta-analysis using the Bucher method enables indirect comparison between '
        'treatments sharing a common comparator(20), provided the transitivity assumption holds: '
        'that the distribution of effect modifiers is sufficiently similar across compared trial '
        'arms. The known demographic and dosing differences between NMN and NR trial populations '
        'represent the principal threat to this assumption.'
    )

    add_para(
        'The objectives of this systematic review were to: (1) characterize structural '
        'heterogeneity of the NMN and NR trial evidence bases through formal transitivity '
        'assessment; (2) conduct pairwise meta-analyses for each precursor versus placebo; '
        '(3) estimate indirect NMN versus NR comparisons via the Bucher method(20) for outcomes '
        'where methodologically feasible; and (4) evaluate certainty of evidence using the '
        'GRADE/CINeMA framework(21). To our knowledge, this is the first study to formally '
        'characterize the evidence gap between NMN and NR and the first indirect comparison '
        'meta-analysis for these precursors.'
    )

    # ================================================================
    # 2. METHODS (~1000 words)
    # ================================================================
    add_heading('2. Methods', level=1)

    add_para(
        'This systematic review and indirect comparison meta-analysis was conducted and reported '
        'in accordance with the Preferred Reporting Items for Systematic Reviews and '
        'Meta-Analyses (PRISMA) 2020 guidelines(17) and the PRISMA extension statement for '
        'network meta-analyses(18). The review was registered in PROSPERO prior to data '
        'screening (March 2026) (registration ID: CRD420261330487). The manuscript title was '
        'subsequently revised to better reflect the primary contribution as an evidence gap '
        'characterization; the registered and submitted titles differ accordingly.'
    )

    add_heading('2.1 Eligibility Criteria', level=2)

    add_para(
        'Studies were eligible if they met the following PICO criteria: (P) adult participants '
        '(>=18 years) of any health status; (I) oral supplementation with NMN or NR at any '
        'dose; (C) placebo comparator; (O) at least one quantitative metabolic outcome '
        '(glycemic markers, lipid profile, hepatic enzymes, blood pressure, anthropometric '
        'measures, or blood NAD+ levels). Only RCTs (parallel or crossover design) were '
        'included. Non-randomized studies, animal or in vitro studies, reviews, protocols, '
        'combination supplements, and pediatric populations were excluded.'
    )

    add_heading('2.2 Information Sources and Search Strategy', level=2)

    add_para(
        'Five electronic databases were searched from January 2018 to May 2025 (the 2018 start '
        'date was selected to coincide with the earliest registered clinical trials of '
        'commercially available NAD+ precursor formulations): PubMed (n=286), Embase (n=375), '
        'Scopus (n=398), Web of Science (n=319), and Cochrane CENTRAL (n=309), yielding 1,687 '
        'records. The search strategy combined terms for the intervention ("nicotinamide '
        'mononucleotide" OR "NMN" OR "nicotinamide riboside" OR "NR" OR "NAD+ precursor") '
        'with study design filters ("randomized" OR "clinical trial" OR "placebo"). The '
        'complete search strategy is provided in Supplementary Table S1. No language '
        'restrictions were applied.'
    )

    add_heading('2.3 Selection Process', level=2)

    add_para(
        'After removal of 562 duplicate records, 1,125 unique records were screened in three '
        'stages: (1) title and abstract screening identified 874 clearly irrelevant records; '
        '(2) 251 records were assessed at full text, of which 217 were excluded (133 upon '
        'full-text review confirming ineligibility apparent from title/abstract criteria, '
        '65 upon re-screening of initially uncertain records, and 19 after detailed full-text '
        'assessment for reasons including non-extractable data, combination interventions, or '
        'out-of-scope populations); (3) 34 reports meeting inclusion criteria were further '
        'assessed, with 19 excluded from qualitative synthesis (trial registrations only, no '
        'extractable data, combination supplements, or duplicate publications). Fifteen studies '
        '(5 NMN, 10 NR) were included in the qualitative synthesis, of which 8 (4 NMN, 4 NR) '
        'contributed data to the quantitative NMA (Figure 1).'
    )

    add_heading('2.4 Data Extraction', level=2)

    add_para(
        'Data were extracted into a standardized form by two independent reviewers. For each '
        'study, we recorded: first author, year, country, design, population, sample size, '
        'dose (mg/day), intervention duration (weeks), participant demographics (age, sex, '
        'BMI), health status, trial registration, and funding source. For each outcome, we '
        'extracted mean change from baseline and standard deviation (SD) for intervention and '
        'control groups. When means and SDs were not directly reported, standard conversions '
        'were applied: SE was converted to SD as SD = SE x sqrt(n); medians and interquartile '
        'ranges were converted using the method of Wan et al.(22). All data points were '
        'independently verified against original publications and supplementary materials. '
        'Inter-rater agreement was substantial to almost perfect across all phases: '
        'title/abstract screening (kappa=0.92), full-text screening (kappa=0.84), RoB 2 '
        'domain-level assessment (kappa=0.75), and data extraction (93.3% agreement; '
        'Supplementary Table S9).'
    )

    add_heading('2.5 Risk of Bias Assessment', level=2)

    add_para(
        'Risk of bias was assessed at the study level using the revised Cochrane Risk of Bias '
        'tool (RoB 2) across five domains: randomization process, deviations from intended '
        'interventions, missing outcome data, measurement of the outcome, and selection of the '
        'reported result(19). Each domain was rated as Low, Some Concerns, or High, with an '
        'overall judgment following RoB 2 algorithms. Detailed domain-level assessments with '
        'justifications are provided in Supplementary Table S4.'
    )

    add_heading('2.6 Data Synthesis and Statistical Analysis', level=2)

    add_para(
        'Pairwise meta-analyses. For each precursor-placebo comparison, random-effects '
        'meta-analyses were conducted using the inverse-variance method with DerSimonian-Laird '
        'estimation of between-study variance (tau-squared)(23). The summary effect measure was '
        'the mean difference (MD) with 95% confidence intervals (CIs). Statistical heterogeneity '
        'was quantified using the I-squared statistic, with thresholds of 25%, 50%, and 75% '
        'indicating low, moderate, and high heterogeneity, respectively.'
    )

    add_para(
        'Network meta-analysis. Because no closed loops of evidence existed (NMN and NR trials '
        'each compared only against placebo), the network geometry was star-shaped. Indirect '
        'comparisons of NMN versus NR were estimated using the Bucher method(20): MD(NMN vs NR) '
        '= MD(NMN vs Placebo) minus MD(NR vs Placebo), with variance equal to the sum of '
        'variances of the two pairwise estimates. Indirect comparison was performed only for '
        'outcomes with compatible measurement definitions and units across NMN and NR arms. '
        'Outcomes with incompatible measurement context or unit scaling were excluded from '
        'indirect estimation.'
    )

    add_para(
        'Sensitivity analyses. Three predefined sensitivity analyses were conducted: '
        '(1) leave-one-out (LOO) analysis for each pairwise comparison, systematically omitting '
        'one study at a time and recording changes in pooled estimates and statistical '
        'significance; (2) LOO analysis for each indirect comparison, propagating LOO pairwise '
        'results through the Bucher formula; (3) exclusion of studies rated High risk of bias '
        'on the overall RoB 2 assessment.'
    )

    add_para(
        'Certainty of evidence. The certainty of evidence for each indirect comparison was '
        'evaluated using the GRADE framework adapted for NMA (CINeMA: Confidence in Network '
        'Meta-Analysis)(21). Six domains were assessed: within-study bias, reporting bias, '
        'indirectness, imprecision, heterogeneity, and incoherence. Each domain was rated as '
        'No Concerns, Some Concerns, or Major Concerns, and certainty was downgraded '
        'accordingly from a starting level of High.'
    )

    add_para(
        'Software. All primary analyses were implemented in Python 3 using scipy (v1.14) for '
        'statistical computations, numpy for numerical operations, and matplotlib for figure '
        'generation. Results were validated using R (v4.5) with the metafor package (v4.8) '
        'using restricted maximum likelihood (REML) estimation, confirming concordance of '
        'pairwise and indirect estimates. Multiplicity adjustment was performed using '
        'Bonferroni and Benjamini-Hochberg corrections across all tested indirectly comparable '
        'outcomes (n=14; NAD+ excluded as non-comparable). The complete analytical pipeline is '
        'available as open-source code.'
    )

    add_para(
        'Publication bias. Formal assessment of publication bias using funnel plots or '
        'Egger\'s regression test was conducted for all pairwise comparisons with two or more '
        'contributing studies (Supplementary Table S10). Formal interpretation was not '
        'performed because the number of studies per comparison was insufficient (k<=4 for '
        'all outcomes; conventional guidance recommends k>=10 for reliable interpretation). '
        'Accordingly, reporting bias was rated as Some Concerns for all outcomes in the '
        'GRADE/CINeMA assessment, based on the predominance of industry funding and the '
        'limited number of available studies.'
    )

    # ================================================================
    # 3. RESULTS (~1800 words)
    # ================================================================
    add_heading('3. Results', level=1)

    add_heading('3.1 Study Selection', level=2)

    add_para(
        'The search identified 1,687 records; after removing 562 duplicates, 1,125 were '
        'screened. Of 251 full-text reports assessed, 217 were excluded. Fifteen studies '
        '(5 NMN, 10 NR) were included in qualitative synthesis, and 8 (4 NMN, 4 NR) '
        'contributed 73 data points spanning 21 metabolic outcomes (15 reported by both '
        'precursors) to quantitative analysis, of which 14 were eligible for indirect '
        'comparison (Figure 1).'
    )
    add_para('[Insert Figure 1 about here]')

    add_heading('3.2 Study Characteristics', level=2)

    add_para(
        'Table 1 summarizes the characteristics of the 15 included studies. The five NMN trials '
        '(Yoshino 2021(2), Huang 2022(3), Igarashi 2022(4), Katayoshi 2023(5), Morifuji '
        '2024(6)) were conducted in the USA (n=1), India (n=1), and Japan (n=3), with doses '
        'of 250-300 mg/day and durations of 8.6-12 weeks. NMN participants were predominantly '
        'from Asian populations with mean BMI ranging from 22.5 to 35.0 kg/m-squared. The ten '
        'NR trials (Dollerup 2018(7), Conze 2019(8), Elhassan 2019(9), Remie 2020(10), '
        'Brakedal 2022(11), Ahmadi 2023(12), Norheim 2024(13), Orr 2024(14), Lin 2025(15), '
        'Bandi 2025(16)) were conducted across eight countries with doses of 500-2,000 mg/day '
        'and durations of 3-12 weeks. NR cohorts were predominantly Western. Total randomized '
        'participants across all 15 studies numbered 740. Study populations ranged from healthy '
        'volunteers to individuals with obesity, prediabetes, chronic kidney disease, '
        'Parkinson\'s disease, COPD, and mild cognitive impairment (Table 1).'
    )
    add_para('[Insert Table 1 about here]')

    add_para(
        'Several systematic differences between NMN and NR trial populations are important for '
        'interpreting the indirect comparisons. NMN trials used a narrow dose range (250-300 '
        'mg/day), whereas NR doses varied nearly fourfold (500-2,000 mg/day), making '
        'dose-equivalence assumptions uncertain. The geographic and ethnic heterogeneity, with '
        'NMN trials concentrated in East Asia and NR trials predominantly in Northern Europe '
        'and North America, introduces potential confounding through differences in dietary '
        'NAD+ intake, genetic variation in NAD+ metabolism enzymes, and baseline '
        'cardiometabolic risk profiles.'
    )

    add_para(
        'To quantify the dosing disparity, we compared molar intakes using the molecular '
        'weights of NMN (334.2 g/mol) and NR chloride (290.7 g/mol, the formulation used in '
        'most NR trials). At 250 mg/day, NMN provides approximately 0.75 mmol/day; at 300 '
        'mg/day, 0.90 mmol/day. NR at 500 mg/day provides 1.72 mmol/day, at 1,000 mg/day '
        '3.44 mmol/day, and at 2,000 mg/day 6.88 mmol/day. Thus NR molar intake ranged from '
        '1.9 to 9.2 times that of NMN across the included trials, with a mean ratio of '
        '5.1-fold. This substantial molar dose asymmetry must be considered when interpreting '
        'comparative efficacy (Table 2).'
    )
    add_para('[Insert Table 2 about here]')

    add_heading('3.3 Transitivity Assessment', level=2)

    add_para(
        'Transitivity requires that the distribution of effect modifiers be sufficiently '
        'similar across NMN and NR trial arms for the placebo node to serve as a valid common '
        'comparator. The NMN and NR trial arms differed substantially on every major effect '
        'modifier (Table 2). The most consequential asymmetry was molar dose: NR trials '
        'delivered a mean of 3.96 mmol/day of precursor versus 0.78 mmol/day for NMN, a '
        '5.1-fold difference on average. Geographic and ethnic separation was systematic, '
        'with NMN trials concentrated in East Asia and NR trials predominantly in Northern '
        'Europe and North America, introducing potential confounding from differences in '
        'dietary NAD+ precursor intake, genetic variation in salvage pathway enzymes (NAMPT, '
        'NRK1, CD38), and baseline cardiometabolic risk profiles. Additionally, 40% of NR '
        'trial participants were drawn from disease-specific populations (Parkinson\'s disease, '
        'chronic kidney disease, COPD, mild cognitive impairment) with no analogous '
        'disease-specific representation in the NMN arm, creating clinical heterogeneity that '
        'further threatens transitivity.'
    )

    add_para(
        'The NAD+ pharmacodynamic endpoint, the most direct marker of precursor '
        'bioavailability, could not be validly compared indirectly because the '
        'NMN-contributing study (Huang 2022) reported blood cellular NAD+/NADH in pmol/mL '
        'while the NR-contributing study (Bandi 2025) reported blood NAD+ in uM, representing '
        'incompatible matrices and scales. This biomarker incompatibility is itself evidence '
        'of the absence of measurement standards in this field.'
    )

    add_para(
        'Collectively, these asymmetries indicate that the transitivity assumption underlying '
        'indirect comparison is substantially threatened in this evidence base. The Very Low '
        'certainty ratings observed across all 14 indirect comparisons are the expected '
        'analytical consequence of this structural inadequacy, not merely a reflection of '
        'small sample sizes.'
    )

    add_heading('3.4 Risk of Bias', level=2)

    add_para(
        'Thirteen studies received an overall RoB 2 rating of "Some Concerns," primarily due '
        'to insufficient randomization detail, industry funding, and metabolic outcomes reported '
        'as secondary endpoints. Two studies were rated "High": Igarashi 2022(4) due to 52% '
        'attrition from a supplement preparation error, and Elhassan 2019(9) due to non-registration '
        'and limited extractable metabolic data. Neither contributed to quantitative synthesis '
        '(Figure 2; Supplementary Table S4).'
    )
    add_para('[Insert Figure 2A about here]')
    add_para('[Insert Figure 2B about here]')

    add_heading('3.5 Pairwise Meta-Analyses', level=2)

    add_para(
        'Random-effects pairwise meta-analyses were conducted for each precursor versus placebo. '
        'Complete numerical results are presented in Supplementary Table S5 with individual '
        'forest plots in Supplementary Figures S1-S5. Key findings are summarized below.'
    )

    add_para(
        'Glycemic outcomes. Fasting blood glucose showed no significant difference for NMN '
        'versus placebo (MD 0.02 mg/dL [95% CI: -4.25, 4.30]; p=0.99; k=4; I-squared=47%) '
        'or NR versus placebo (MD -2.58 [-6.60, 1.43]; p=0.21; k=2; I-squared=0%). HbA1c '
        'was unchanged with both NMN (MD -0.01% [-0.14, 0.12]; p=0.91; k=2) and NR '
        '(MD -0.20% [-0.48, 0.08]; p=0.16; k=1). Given that HbA1c reflects average glycemia '
        'over 8-12 weeks, the short trial durations may have been insufficient to detect '
        'meaningful shifts. HOMA-IR showed a non-significant reduction with NMN in a single '
        'study (MD -0.60 [-1.40, 0.20]; p=0.14; k=1) and NR (MD -0.19 [-0.90, 0.51]; '
        'p=0.59; k=1). Fasting insulin showed no significant effect for either precursor '
        '(NMN: MD -3.29 uIU/mL [-8.28, 1.69]; p=0.20; k=2; I-squared=0%; NR: MD -0.42 '
        '[-2.65, 1.82]; p=0.71; k=1). Note that 17 of 36 pairwise comparisons across all '
        'outcomes are based on a single study (k=1) and therefore represent individual trial '
        'estimates rather than pooled meta-analytic results.'
    )

    add_para(
        'Lipid profile. No significant effects were observed for total cholesterol '
        '(NMN: MD -4.16 mg/dL [-17.03, 8.71]; p=0.53; k=2; I-squared=0%; NR: MD -7.92 '
        '[-19.87, 4.02]; p=0.19; k=3), LDL cholesterol (NMN: MD -1.52 [-11.68, 8.64]; '
        'p=0.77; k=3; NR: MD -5.07 [-14.71, 4.57]; p=0.30; k=3), HDL cholesterol '
        '(NMN: MD -0.15 [-5.73, 5.44]; p=0.96; k=3; NR: MD -2.19 [-7.59, 3.22]; p=0.43; '
        'k=3), or triglycerides (NMN: MD -17.78 [-36.53, 0.98]; p=0.063; k=4; NR: MD 4.17 '
        '[-25.38, 33.71]; p=0.78; k=3; I-squared=20%). The NMN triglyceride estimate was '
        'numerically the largest in magnitude but remained non-significant (p=0.063) and '
        'should be interpreted cautiously given imprecision. Wide confidence intervals '
        'across all lipid comparisons reflect the small number of contributing studies.'
    )

    add_para(
        'Hepatic enzymes. ALT was not significantly changed by NMN (MD -1.13 U/L '
        '[-5.13, 2.87]; p=0.58; k=3; I-squared=66%) or NR (MD -3.26 [-7.77, 1.25]; '
        'p=0.16; k=2). AST was similarly non-significant for either precursor '
        '(NMN: MD -0.30 [-1.99, 1.38]; p=0.72; k=3; I-squared=18%; NR: MD -1.40 '
        '[-4.34, 1.54]; p=0.35; k=1). The moderate heterogeneity for ALT NMN '
        '(I-squared=66%) suggests variability in hepatic response across studies.'
    )

    add_para(
        'Blood pressure. SBP showed a non-significant reduction with NMN '
        '(MD -3.50 mmHg [-7.85, 0.85]; p=0.11; k=4; I-squared=0%) but not with NR '
        '(MD 2.00 [-5.84, 9.84]; p=0.62; k=1). DBP was similarly non-significant for '
        'NMN (MD -2.55 [-5.60, 0.51]; p=0.10; k=4) and NR (MD 0.80 [-4.47, 6.07]; '
        'p=0.77; k=1). SBP and DBP point estimates were directionally negative for NMN '
        'but both remained non-significant with confidence intervals crossing the null; '
        'these trends require prospective confirmation in adequately powered trials.'
    )

    add_para(
        'Anthropometric measures. Neither BMI (NMN: MD -0.04 kg/m-squared [-1.29, 1.21]; '
        'p=0.95; k=3; NR: MD -0.90 [-3.41, 1.61]; p=0.48; k=1) nor body weight '
        '(NMN: MD 0.61 kg [-3.80, 5.02]; p=0.79; k=3; NR: MD -0.30 [-3.22, 2.62]; '
        'p=0.84; k=1) was significantly altered by either precursor. These null findings '
        'are unsurprising given that trial durations were too short for meaningful body '
        'composition changes, and none of the trials were designed with weight loss as '
        'a primary endpoint.'
    )

    add_para(
        'Blood NAD+. Blood NAD+ is the most direct pharmacodynamic marker of precursor '
        'supplementation. However, the two contributing studies measured different constructs '
        'and scales: Huang 2022 reported blood cellular NAD+/NADH (pmol/mL), whereas Bandi '
        '2025 reported blood NAD+ (uM). Because these are not directly harmonizable for a '
        'valid indirect contrast, NAD+ was retained as separate pairwise evidence within '
        'each precursor but excluded from NMN-versus-NR indirect comparison.'
    )

    add_heading('3.6 Network Meta-Analysis: Indirect Comparisons', level=2)

    add_para(
        'The network geometry was star-shaped, with NMN and NR connected only through the '
        'common placebo node (Figure 3). Table 3 presents Bucher indirect comparison results '
        'for 14 metabolically comparable outcomes (NAD+ excluded for incompatibility of '
        'measurement context and units).'
    )
    add_para('[Insert Figure 3 about here]')
    add_para('[Insert Table 3 about here]')

    add_para(
        'No indirect comparison showed a statistically significant difference between NMN '
        'and NR across the 14 comparable metabolic outcomes (all p>0.10). Confidence '
        'intervals for most outcomes were wide and crossed the null, reflecting substantial '
        'imprecision driven by small sample sizes and sparse networks. Given 14 simultaneous '
        'indirect comparisons and generally low information size, these estimates should be '
        'interpreted as hypothesis-generating rather than confirmatory (Figure 4).'
    )
    add_para('[Insert Figure 4 about here]')

    add_heading('3.7 Certainty of Evidence', level=2)

    add_para(
        'All 14 indirect comparisons were rated Very Low certainty by GRADE/CINeMA(21) '
        '(Figure 5; Supplementary Table S8). Downgrading was driven by within-study bias '
        '(Some Concerns in all comparisons, reflecting the preponderance of industry funding '
        'and secondary reporting of metabolic outcomes), reporting bias (Some Concerns due '
        'to insufficient studies for funnel plot assessment), indirectness (Some Concerns '
        'due to population heterogeneity, as NMN trials enrolled predominantly Asian '
        'participants while NR trials enrolled predominantly Western participants, and dose '
        'ranges differed substantially: 250 mg/day for NMN vs 500-2,000 mg/day for NR), '
        'imprecision (Some to Major Concerns), and heterogeneity (No to Some Concerns). '
        'Incoherence was Not Applicable for all outcomes because the star-shaped network '
        'contained no closed loops.'
    )
    add_para('[Insert Figure 5 about here]')

    add_heading('3.8 Sensitivity Analyses', level=2)

    add_para(
        'Leave-one-out analyses were performed for all 56 pairwise iterations across all '
        'outcomes and comparisons. No pairwise significance changes were identified in any '
        'leave-one-out iteration, indicating that no single study exerts disproportionate '
        'influence on any pairwise estimate.'
    )

    add_para(
        'Leave-one-out analysis of the 13 indirect comparisons amenable to LOO (56 total '
        'iterations; HOMA-IR excluded because both arms comprised a single study; '
        'Supplementary Table S7) revealed one significance change: the triglyceride indirect '
        'comparison became nominally significant (p=0.046, uncorrected) upon exclusion of '
        'Conze 2019, though this did not survive Bonferroni correction (threshold p<0.0036), '
        'with the effect driven by the remaining NR studies. All other indirect comparisons '
        'remained non-significant regardless of which study was excluded (Figure 6).'
    )
    add_para('[Insert Figure 6 about here]')

    add_para(
        'The high-RoB exclusion analysis had no effect on any result, as neither high-RoB '
        'study (Igarashi 2022(4), Elhassan 2019(9)) contributed data to the quantitative '
        'meta-analysis. Igarashi 2022 was excluded due to 52% attrition from a supplement '
        'preparation error, and Elhassan 2019 was excluded due to non-registration and '
        'insufficient extractable metabolic data.'
    )

    add_heading('3.9 Exploratory Subgroup Observations', level=2)

    add_para(
        'Although the small number of studies precluded formal subgroup meta-analysis, '
        'a narrative comparison provides informative context. Among NMN trials, Yoshino '
        '2021(2) enrolled prediabetic obese women (BMI>30) and showed the largest '
        'insulin-sensitivity improvements, although results did not reach significance '
        'as single-study estimates. Katayoshi 2023(5) and Morifuji 2024(6), which '
        'enrolled healthy volunteers, showed no glycemic improvements. Among NR trials, '
        'Dollerup 2018(7) enrolled obese men and reported numerical but non-significant '
        'improvements in hepatic lipid content.'
    )

    add_para(
        'This pattern suggests that baseline metabolic status may be an important effect '
        'modifier: populations with established insulin resistance or dyslipidemia may '
        'benefit more from NAD+ repletion than metabolically healthy individuals. The '
        'floor effect in healthy populations may obscure meaningful treatment effects. '
        'Future trials should stratify by metabolic status to test this hypothesis.'
    )

    # ================================================================
    # 4. DISCUSSION (~1200 words)
    # ================================================================
    add_heading('4. Discussion', level=1)

    add_heading('4.1 Summary of Findings', level=2)

    add_para(
        'This systematic review provides the first formal characterization of the evidence gap '
        'between NMN and NR for metabolic outcomes, and the first indirect comparison '
        'meta-analysis to demonstrate why that gap cannot yet be reliably closed. The primary '
        'finding is structural: the two evidence bases are systematically incompatible for '
        'reliable indirect comparison, with a 1.9- to 9.2-fold molar dose asymmetry, systematic '
        'ethnic and geographic separation, and incompatible NAD+ biomarker definitions. None of '
        'the 14 indirect comparisons reached statistical significance, and all were rated Very '
        'Low certainty(21).'
    )

    add_heading('4.2 Interpretation of the NAD+ Finding', level=2)

    add_para(
        'The NAD+ findings require careful interpretation because the available NMN and NR '
        'studies did not report directly comparable biomarkers. Huang 2022 reported blood '
        'cellular NAD+/NADH in pmol/mL, while Bandi 2025 reported blood NAD+ in uM. This '
        'methodological mismatch precludes a valid indirect NMN-versus-NR contrast for NAD+ '
        'and reinforces the need for future trials to use harmonized assays, matrix '
        'definitions, and unit reporting.'
    )

    add_para(
        'Mechanistically, first-pass metabolism remains relevant as within-arm '
        'pharmacodynamic context rather than as evidence for an indirect NMN-versus-NR '
        'effect. A substantial fraction of orally ingested NR is cleaved to nicotinamide '
        'by purine nucleoside phosphorylase in the gut lumen and liver before reaching the '
        'systemic circulation(30). Despite this hepatic catabolism, several NR trials still '
        'showed substantial pairwise NAD+ increases, possibly because higher absolute NR '
        'doses (up to 2,000 mg/day) compensated for first-pass losses. NMN may also be '
        'subject to gastrointestinal degradation, but the extent of this process in humans '
        'remains insufficiently characterized. From a clinical perspective, future trials '
        'should prioritize harmonized NAD+ assay matrices, unit definitions, and reporting '
        'standards so that pharmacodynamic findings can be compared across precursors and '
        'linked to patient-relevant metabolic outcomes.'
    )

    add_heading('4.3 Context with Prior Reviews', level=2)

    add_para(
        'Our pairwise findings are broadly consistent with recent systematic reviews. Zheng '
        'et al.(24) conducted the most comprehensive NMN-focused meta-analysis to date, '
        'pooling 9 RCTs (n=412) and reporting non-significant effects on fasting glucose, '
        'HbA1c, HOMA-IR, lipid profile, and blood pressure. Our NMN pairwise estimates are '
        'directionally concordant with theirs, though minor numerical differences exist due '
        'to differences in study inclusion and data extraction; notably, four NMN trials '
        'included by Zheng et al. did not meet the present review\'s eligibility and '
        'harmonized extraction criteria for quantitative synthesis. Alegre and Pastore(25) '
        'provided a narrative review of both NMN and NR, concluding that while preclinical '
        'evidence was promising, clinical translation remained limited; they specifically '
        'highlighted the absence of head-to-head trials as a critical gap. Nascimento and '
        'Nogueira-de-Almeida(26) focused on NR in cardiometabolic disease and found that, '
        'despite NR\'s consistent ability to raise circulating NAD+, downstream metabolic '
        'benefits were not reliably observed.'
    )

    add_para(
        'This review extends prior work by formally comparing NMN against NR through Bucher '
        'indirect comparison(20) and, critically, by providing the first systematic diagnosis '
        'of why that comparison produces uninformative results. Prior reviews synthesized '
        'each precursor in isolation against placebo; none characterized the structural '
        'compatibility of the two evidence bases or quantified the barriers to reliable '
        'indirect inference. The finding that all 14 indirect comparisons returned Very Low '
        'certainty is not a failure of the analysis; it is the informative result. A '
        'well-designed indirect comparison applied to a structurally incompatible evidence '
        'base should return Very Low certainty, and documenting this formally is the '
        'necessary precondition for designing trials that can do better.'
    )

    add_heading('4.4 Explanations for Null Metabolic Findings', level=2)

    add_para(
        'Several factors likely contribute to the absence of detectable metabolic differences '
        'between NMN and NR. The included trials had small sample sizes (median n approximately '
        '40), providing insufficient statistical power; a post hoc illustrative calculation '
        '(two-sided alpha=0.05, 80% power, target difference=5 mg/dL, assumed common SD '
        'approximately 16 mg/dL based on observed trial-level FBG variability) indicates that '
        'approximately 80 participants per arm would be required in a direct two-arm comparison, '
        'far exceeding the effective sample sizes here. Trial durations were short (median 8-12 '
        'weeks), likely too brief for sirtuin-dependent metabolic reprogramming to produce '
        'measurable changes in clinical biomarkers(33). Most study populations were metabolically '
        'healthy or mildly impaired, creating a floor effect. Numerical but non-significant '
        'improvements in HOMA-IR were observed with both NMN (MD -0.60; Huang 2022(3)) and '
        'NR (MD -0.19; Dollerup 2018(7)), suggesting that metabolically at-risk populations '
        'may show greater responsiveness to NAD+ repletion.'
    )

    add_para(
        'The dose asymmetry between NMN and NR trials warrants particular attention. On a '
        'mass basis, NMN was tested at 250-300 mg/day while NR was tested at 500-2,000 '
        'mg/day. Our molar dose analysis (Section 3.2) revealed that NR molar intake was '
        '1.9-9.2-fold greater than NMN across included trials. This asymmetry is a major '
        'threat to the transitivity assumption and limits interpretability of indirect '
        'contrasts. Future head-to-head trials should employ equimolar dosing to enable '
        'pharmacologically meaningful comparison.'
    )

    add_para(
        'Finally, circulating NAD+ may be an imperfect surrogate for tissue-level NAD+ in '
        'metabolically relevant organs. The distribution of NAD+ to skeletal muscle, liver, '
        'and adipose tissue may differ between precursors even when blood levels diverge, '
        'and tissue-specific measurements using muscle biopsy or advanced imaging techniques '
        'would be needed to resolve this question.'
    )

    add_heading('4.5 Strengths and Limitations', level=2)

    add_para(
        'This study has several notable strengths. It is the first indirect comparison '
        'meta-analysis and formal transitivity assessment of NMN versus NR for metabolic '
        'outcomes. We searched five databases with broad inclusion criteria, extracted 21 '
        'distinct metabolic outcomes with 73 data points, performed GRADE/CINeMA(21) '
        'certainty assessment for all 14 indirectly comparable outcomes, and conducted '
        'multiple predefined sensitivity analyses. All data extraction was independently '
        'performed by two reviewers with substantial inter-rater agreement '
        '(kappa=0.75-0.92; Supplementary Table S9), and primary Python-based analyses '
        'were independently validated using R/metafor with REML estimation. The complete '
        'dataset and analytical code are publicly available.'
    )

    add_para(
        'Several important limitations must be acknowledged. The star-shaped network '
        'geometry precludes assessment of incoherence (consistency), a core NMA diagnostic. '
        'Seventeen of 36 pairwise comparisons (47%) are based on a single study (k=1), '
        'representing single-trial estimates rather than pooled results. NAD+ could not be '
        'compared indirectly due to non-comparable biomarker definitions. For crossover '
        'trials, between-person variance was used because within-person correlation '
        'parameters were unavailable, potentially overestimating variance.'
    )

    add_para(
        'Population heterogeneity between NMN and NR trials represents the most serious '
        'threat to transitivity. NMN trials were predominantly conducted in Asian '
        'populations, whereas NR trials were predominantly Western. Genetic polymorphisms '
        'in NAD+ metabolic enzymes (NAMPT, NRK1, CD38) differ across ethnic groups and '
        'could differentially affect precursor efficacy. Dietary patterns also vary '
        'systematically, potentially affecting baseline NAD+ status. Dose asymmetry '
        '(250 mg NMN vs 500-2,000 mg NR) introduces further heterogeneity unresolvable '
        'without pharmacokinetic equivalence studies. Publication bias could not be '
        'formally assessed (k<=4 for all comparisons; recommended k>=10). Industry '
        'funding raises reporting bias concerns, particularly for metabolic outcomes that '
        'were often secondary endpoints. The Very Low certainty of all 14 outcomes means '
        'true effects could be substantially different from estimates.'
    )

    add_heading('4.6 Implications for Practice and Research', level=2)

    add_para(
        'For clinical practice, current evidence does not support choosing NMN over NR or vice '
        'versa for any metabolic indication. Marketing claims of superiority for either precursor '
        'are not supported by the evidence base. This finding is directly relevant to public '
        'health nutrition policy, given the widespread marketing of these supplements to consumers '
        'worldwide.'
    )

    add_para(
        'For future research, the structural barriers identified are quantifiable and actionable. '
        'A definitive head-to-head trial should: (1) use equimolar dosing (approximately 1,150 mg '
        'NMN/day equivalent to 1,000 mg NR/day); (2) harmonize NAD+ assessment using whole-blood '
        'enzymatic cycling assays reported in umol/L; (3) enrol metabolically at-risk populations '
        '(HOMA-IR >=2.5 or fasting glucose 100-125 mg/dL); (4) run for a minimum of 24 weeks; '
        '(5) include more than 100 participants per arm; and (6) use funding independent of '
        'supplement manufacturers. None of the 15 trials in this review met more than two of '
        'these criteria.'
    )

    # ================================================================
    # 5. CONCLUSION (~140 words)
    # ================================================================
    add_heading('5. Conclusion', level=1)

    add_para(
        'In this systematic review and indirect comparison meta-analysis, we provide the first '
        'formal characterization of the evidence gap between nicotinamide mononucleotide and '
        'nicotinamide riboside for metabolic outcomes. The primary finding is structural: the '
        'NMN and NR trial evidence bases are systematically incompatible for reliable indirect '
        'comparison, differing by 1.9- to 9.2-fold on a molar dose basis, separated by '
        'geography and ethnicity in ways that threaten transitivity, and using incompatible '
        'NAD+ biomarker definitions that preclude pharmacodynamic comparison. The consequence '
        'of these structural problems is precisely what the analysis found: 14 indirect '
        'comparisons across clinically relevant metabolic outcomes, all non-significant, all '
        'rated Very Low certainty. These results do not indicate that NMN and NR are '
        'equivalent; they indicate that the current evidence base cannot distinguish them. '
        'Closing this gap requires head-to-head trials with equimolar dosing (approximately '
        '1,150 mg NMN equivalent to 1,000 mg NR), harmonized NAD+ assay standards, '
        'metabolically at-risk populations, and independent funding. Until such trials exist, '
        'evidence-based guidance on choosing between NMN and NR for metabolic health remains '
        'impossible.'
    )

    # ================================================================
    # ACKNOWLEDGEMENTS
    # ================================================================
    add_heading('Acknowledgements', level=1)
    add_para('None.')

    # ================================================================
    # SUPPLEMENTARY MATERIALS
    # ================================================================
    add_heading('Supplementary Materials', level=1)
    add_para(
        'The following supplementary materials are available online:\n'
        'Supplementary Table S1. Complete search strategy.\n'
        'Supplementary Table S2. Detailed study characteristics.\n'
        'Supplementary Table S3. Full data extraction table.\n'
        'Supplementary Table S4. RoB 2 domain-level assessments.\n'
        'Supplementary Table S5. Complete pairwise meta-analysis results.\n'
        'Supplementary Table S6. Complete NMA results.\n'
        'Supplementary Table S7. Leave-one-out sensitivity analysis summary.\n'
        'Supplementary Table S8. GRADE/CINeMA certainty assessment.\n'
        'Supplementary Table S9. Inter-rater agreement.\n'
        'Supplementary Table S10. Publication bias assessment.\n'
        'Supplementary Figures S1-S5. Pairwise forest plots by outcome category.\n'
        'Supplementary Figure S6. Outcome reporting matrix.\n'
        'Supplementary Figure S7. Representative leave-one-out forest plots.\n'
        'PRISMA 2020 Checklist.'
    )

    # ================================================================
    # FIGURE LEGENDS (grouped at end per PHN)
    # ================================================================
    add_heading('Figure Legends', level=1)

    add_para(
        'Figure 1. PRISMA 2020 flow diagram showing study identification, screening, '
        'eligibility, and inclusion.'
    )
    add_para(
        'Figure 2A. Risk of bias traffic-light plot showing domain-level RoB 2 judgments '
        'for each included study.'
    )
    add_para(
        'Figure 2B. Risk of bias summary bar chart showing the distribution of Low, Some '
        'Concerns, and High ratings across all studies for each domain.'
    )
    add_para(
        'Figure 3. Network diagram showing the star-shaped geometry with NMN, NR, and Placebo '
        'nodes. Edge thickness is proportional to the number of studies.'
    )
    add_para(
        'Figure 4. Forest plot of indirect comparisons (NMN vs NR) for 14 metabolically '
        'comparable outcomes.'
    )
    add_para(
        'Figure 5. GRADE/CINeMA certainty-of-evidence heatmap for 14 indirectly comparable '
        'NMN vs NR outcomes.'
    )
    add_para(
        'Figure 6. Leave-one-out sensitivity summary for all indirect NMN vs NR comparisons.'
    )

    # ================================================================
    # REFERENCES
    # ================================================================
    add_heading('References', level=1)

    refs = [
        'Yoshino J, Baur JA, Imai S-I. NAD+ intermediates: The biology and therapeutic potential of NMN and NR. Cell Metab. 2018;27(3):513-528.',
        'Yoshino M, Yoshino J, Kayser BD, et al. Nicotinamide mononucleotide increases muscle insulin sensitivity in prediabetic women. Science. 2021;372(6547):1224-1229.',
        'Huang H. A multicentre, randomised, double-blind, parallel design, placebo controlled study to evaluate the efficacy and safety of Uthever (NMN supplement). Front Aging. 2022;3:851698.',
        'Igarashi M, Nakagawa-Nagahama Y, Miura M, et al. Chronic nicotinamide mononucleotide supplementation elevates blood nicotinamide adenine dinucleotide levels and alters muscle function in healthy older men. npj Aging. 2022;8:5.',
        'Katayoshi T, Uehata S, Nakashima N, et al. Nicotinamide mononucleotide (NMN) intake increases plasma NMN and insulin levels in healthy subjects. Endocr J. 2023;70(10):1003-1011.',
        'Morifuji M, Higashi S, Ebihara S, Nagata K. A randomized, double-blind, placebo-controlled trial of the effects of nicotinamide mononucleotide supplementation on biomarkers of aging in older Japanese adults. Chin Med J. 2024;137(19):2365-2374.',
        'Dollerup OL, Christensen B, Svart M, et al. A randomized placebo-controlled clinical trial of nicotinamide riboside in obese men: safety, insulin-sensitivity, and lipid-mobilizing effects. Am J Clin Nutr. 2018;108(2):343-353.',
        'Conze D, Brenner C, Kruger CL. Safety and metabolism of long-term administration of NIAGEN (nicotinamide riboside chloride) in a randomized, double-blind, placebo-controlled clinical trial of healthy overweight adults. Sci Rep. 2019;9:9772.',
        'Elhassan YS, Kluckova K, Fletcher RS, et al. Nicotinamide riboside augments the aged human skeletal muscle NAD+ metabolome and induces transcriptomic and anti-inflammatory signatures. Cell Rep. 2019;28(7):1717-1728.',
        'Remie CME, Roumans KHM, Moonen MPB, et al. Nicotinamide riboside supplementation alters body composition and skeletal muscle acetylcarnitine concentrations in healthy obese humans. Am J Clin Nutr. 2020;112(2):413-426.',
        'Brakedal B, Dolle C, Riber F, et al. The NADPARK study: A randomized phase I trial of nicotinamide riboside supplementation in Parkinson\'s disease. Cell Metab. 2022;34(3):396-407.',
        'Ahmadi A, Begue G, Valencia AP, et al. Randomized crossover clinical trial of niacin and nicotinamide riboside on NAD+ in people with moderate-to-severe CKD. Kidney Int Rep. 2023;8:1055-1068.',
        'Norheim KL, Ben Ezra M, Heckenbach I, et al. Effect of nicotinamide riboside on airway inflammation in COPD: a randomized, placebo-controlled trial. Nat Aging. 2024;4:1772-1781.',
        'Orr ME, Kotkowski E, Rabinovici GD, et al. Nicotinamide riboside in people with mild cognitive impairment (NR-MCI): A pilot study. Alzheimers Dement (N Y). 2024;10:e12444.',
        'Lin J, King AC. Nicotinamide riboside and exercise to augment vascular and brain health in hypertensive older adults: pilot results from the NREX trial. Nutrients. 2025;17(1):112.',
        'Bandi R, Gupta R, Koppala D, et al. A randomized, double-blind, placebo-controlled clinical trial assessing efficacy and safety of LN22199 and nicotinamide riboside on biomarkers of aging. J Gerontol A. 2025;80(2):glae260.',
        'Page MJ, McKenzie JE, Bossuyt PM, et al. The PRISMA 2020 statement: an updated guideline for reporting systematic reviews. BMJ. 2021;372:n71.',
        'Hutton B, Salanti G, Caldwell DM, et al. The PRISMA extension statement for reporting of systematic reviews incorporating network meta-analyses. Ann Intern Med. 2015;162(11):777-784.',
        'Sterne JAC, Savovic J, Page MJ, et al. RoB 2: a revised tool for assessing risk of bias in randomised trials. BMJ. 2019;366:l4898.',
        'Bucher HC, Guyatt GH, Griffith LE, Walter SD. The results of direct and indirect treatment comparisons in meta-analysis of randomized controlled trials. J Clin Epidemiol. 1997;50(6):683-691.',
        'Nikolakopoulou A, Higgins JPT, Papakonstantinou T, et al. CINeMA: An approach for assessing confidence in the results of a network meta-analysis. PLoS Med. 2020;17(4):e1003082.',
        'Wan X, Wang W, Liu J, Tong T. Estimating the sample mean and standard deviation from the sample size, median, range and/or interquartile range. BMC Med Res Methodol. 2014;14:135.',
        'DerSimonian R, Laird N. Meta-analysis in clinical trials. Control Clin Trials. 1986;7(3):177-188.',
        'Zheng Y, Guo Y, Chen Y, et al. Nicotinamide mononucleotide supplementation and cardiometabolic health: a systematic review and meta-analysis of randomized controlled trials. Nutrients. 2024;16(15):2526.',
        'Alegre GFS, Pastore GM. NAD+ precursors nicotinamide mononucleotide (NMN) and nicotinamide riboside (NR): Potential dietary contribution to health. Curr Nutr Rep. 2024;13:261-282.',
        'Nascimento GP, Nogueira-de-Almeida CA. Nicotinamide riboside supplementation in adults with cardiometabolic disease: a systematic review. Eur J Clin Nutr. 2024;78:815-826.',
        'Grozio A, Mills KF, Yoshino J, et al. Slc12a8 is a nicotinamide mononucleotide transporter. Nat Metab. 2019;1(1):47-57.',
        'Canto C, Houtkooper RH, Pirinen E, et al. The NAD+ precursor nicotinamide riboside enhances oxidative metabolism and protects against high-fat diet-induced obesity. Cell Metab. 2012;15(6):838-847.',
        'Mills KF, Yoshida S, Stein LR, et al. Long-term administration of nicotinamide mononucleotide mitigates age-associated physiological decline in mice. Cell Metab. 2016;24(6):795-806.',
        'Trammell SAJ, Schmidt MS, Weidemann BJ, et al. Nicotinamide riboside is uniquely and orally bioavailable in mice and humans. Nat Commun. 2016;7:12948.',
        'Verdin E. NAD+ in aging, metabolism, and neurodegeneration. Science. 2015;350(6265):1208-1213.',
        'Camacho-Pereira J, Tarrago MG, Chini CCS, et al. CD38 dictates age-related NAD+ decline and mitochondrial dysfunction through an SIRT3-dependent mechanism. Cell Metab. 2016;23(6):1127-1139.',
        'Imai S, Guarente L. NAD+ and sirtuins in aging and disease. Trends Cell Biol. 2014;24(8):464-471.',
    ]

    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.text = '(%d) %s' % (i, ref)

    # ================================================================
    # TABLES (built from CSV data)
    # ================================================================
    doc.add_page_break()
    add_heading('Tables', level=1)

    # ------ Table 1: Study characteristics ------
    add_para('Table 1. Characteristics of included studies (n=15).')
    t1_headers = ['Study', 'Precursor', 'Country', 'Design', 'N', 'Dose\n(mg/d)',
                  'Duration\n(wk)', 'Health status']
    import csv
    t1_data = []
    with open('results/tables/supp_table_S2_study_characteristics.csv') as f:
        reader = csv.DictReader(f)
        for row in reader:
            t1_data.append([
                row['Study'], row['Precursor'], row['Country'],
                row['Design'].replace('Parallel ', '').replace(' RCT', ''),
                row['N_randomized'], row['Dose_mg_day'].replace('mg', ''),
                row['Duration_weeks'], row['Health_status']
            ])
    table1 = doc.add_table(rows=1 + len(t1_data), cols=len(t1_headers))
    table1.style = 'Table Grid'
    for j, h in enumerate(t1_headers):
        cell = table1.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, row_data in enumerate(t1_data):
        for j, val in enumerate(row_data):
            cell = table1.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    add_para('Abbreviations: NMN, nicotinamide mononucleotide; NR, nicotinamide riboside; '
             'wk, weeks; N, number randomized.')
    doc.add_paragraph()

    # ------ Table 2: Transitivity assessment ------
    add_para('Table 2. Transitivity assessment: comparison of effect modifier distributions '
             'across NMN and NR trial arms.')
    t2_headers = ['Domain', 'NMN arm', 'NR arm', 'Assessment']
    t2_data = []
    with open('results/tables/transitivity_assessment.csv') as f:
        reader = csv.DictReader(f)
        for row in reader:
            t2_data.append([row['Domain'], row['NMN_arm'], row['NR_arm'],
                            row['Assessment']])
    table2 = doc.add_table(rows=1 + len(t2_data), cols=len(t2_headers))
    table2.style = 'Table Grid'
    for j, h in enumerate(t2_headers):
        cell = table2.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, row_data in enumerate(t2_data):
        for j, val in enumerate(row_data):
            cell = table2.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

    # ------ Table 3: Indirect comparisons (NMN vs NR) ------
    add_para('Table 3. Indirect comparisons of NMN versus NR via Bucher method '
             '(14 outcomes; NAD+ excluded as non-comparable).')
    t3_headers = ['Outcome', 'MD', '95% CI', 'P', 'k (NMN+NR)']
    t3_data = []
    with open('results/tables/nma_results.csv') as f:
        reader = csv.DictReader(f)
        for row in reader:
            if row['type'] == 'indirect':
                md = float(row['MD'])
                lo = float(row['lower_CI'])
                hi = float(row['upper_CI'])
                pv = float(row['p_value'])
                outcome_label = row['outcome'].replace('_', ' ')
                if outcome_label == 'body weight':
                    outcome_label = 'Body weight'
                elif outcome_label == 'fasting insulin':
                    outcome_label = 'Fasting insulin'
                else:
                    outcome_label = outcome_label.upper() if len(outcome_label) <= 6 else outcome_label.capitalize()
                    if outcome_label == 'Homa-ir':
                        outcome_label = 'HOMA-IR'
                    elif outcome_label == 'Hba1c':
                        outcome_label = 'HbA1c'
                ci_str = '[%.2f, %.2f]' % (lo, hi)
                p_str = '%.4f' % pv if pv >= 0.001 else '<0.001'
                t3_data.append([outcome_label, '%.2f' % md, ci_str, p_str,
                                row['k']])
    table3 = doc.add_table(rows=1 + len(t3_data), cols=len(t3_headers))
    table3.style = 'Table Grid'
    for j, h in enumerate(t3_headers):
        cell = table3.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, row_data in enumerate(t3_data):
        for j, val in enumerate(row_data):
            cell = table3.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    add_para('Abbreviations: MD, mean difference; CI, confidence interval; k, number of '
             'studies contributing to each arm (NMN+NR).')

    # Save
    path = os.path.join(BASE, 'NMN_NR_Systematic_Review_PHN.docx')
    doc.save(path)
    print("Created: NMN_NR_Systematic_Review_PHN.docx")

    # ---- Word count verification ----
    body_words = 0
    in_body = False
    exclude_sections = {'Abstract', 'References', 'Figure Legends', 'Tables',
                        'Acknowledgements', 'Supplementary Materials'}
    current_section = None

    for p in doc.paragraphs:
        text = p.text.strip()
        if p.style and 'Heading 1' in p.style.name:
            current_section = text.split('. ', 1)[-1] if '. ' in text else text
            # Check if this is the start of a body section
            if text.startswith('1.') or text.startswith('2.') or text.startswith('3.') or \
               text.startswith('4.') or text.startswith('5.'):
                in_body = True
            elif current_section in exclude_sections:
                in_body = False
        elif in_body and text:
            if text.startswith('[Insert '):
                continue
            body_words += len(text.split())

    print("Body text word count (Intro through Conclusion): %d" % body_words)
    if body_words > 5000:
        print("WARNING: Over 5000 word limit by %d words!" % (body_words - 5000))
    else:
        print("Under limit by %d words." % (5000 - body_words))


# ============================================================
# RUN
# ============================================================
if __name__ == '__main__':
    create_title_page()
    create_manuscript()
    print("\nAll PHN submission files created in: %s/" % BASE)
    print("\nReminder: Copy Tables 1, 2, 3 from the original manuscript into the new docx.")
    print("Reminder: Upload PRISMA 2020 checklist as supplementary material.")
    print("Reminder: Figures uploaded as separate files (TIFF, EPS, or PDF preferred).")
